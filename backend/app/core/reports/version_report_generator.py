from pathlib import Path
from typing import List
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.config import settings


class VersionReportGenerator:
    """Generate version history reports"""
    
    def __init__(self):
        self.doc = None
    
    def generate_version_history_report(self, stage, history: List, project_name: str) -> str:
        """
        Generate complete version history report
        
        Args:
            stage: Stage object
            history: List of VersionHistory objects
            project_name: Name of the project
            
        Returns:
            Path to generated DOCX file
        """
        self.doc = Document()
        
        # Title
        title = self.doc.add_heading(f'Version History Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Project and Stage Info
        self.doc.add_heading('Stage Information', level=1)
        info_table = self.doc.add_table(rows=5, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        info_data = [
            ('Project', project_name),
            ('Stage Number', str(stage.stage_number)),
            ('Stage Name', stage.stage_name),
            ('Current Version', stage.version_number),
            ('Total Versions', str(len(history)))
        ]
        
        for i, (key, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = key
            info_table.rows[i].cells[1].text = value
        
        self.doc.add_paragraph()
        
        # Version History Table
        self.doc.add_heading('Version History', level=1)
        
        if history:
            # Create table with headers
            history_table = self.doc.add_table(rows=1, cols=4)
            history_table.style = 'Light Grid Accent 1'
            
            # Headers
            headers = history_table.rows[0].cells
            headers[0].text = 'Version'
            headers[1].text = 'Action'
            headers[2].text = 'Timestamp'
            headers[3].text = 'Details'
            
            # Make headers bold
            for cell in headers:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Add version data
            for v in history:
                row = history_table.add_row().cells
                row[0].text = v.version_number or 'N/A'
                
                metadata = v.version_metadata or {}
                action = metadata.get('action', 'Unknown')
                row[1].text = action
                
                row[2].text = v.timestamp.strftime('%Y-%m-%d %H:%M:%S') if v.timestamp else 'N/A'
                
                # Details
                details = []
                if metadata.get('description'):
                    details.append(metadata['description'])
                if metadata.get('validation_status'):
                    details.append(f"Status: {metadata['validation_status']}")
                if metadata.get('passed') is not None:
                    details.append(f"Passed: {'Yes' if metadata['passed'] else 'No'}")
                
                row[3].text = '\n'.join(details) if details else 'No details'
        else:
            self.doc.add_paragraph('No version history available.')
        
        self.doc.add_paragraph()
        
        # Stage Logic History
        self.doc.add_heading('Stage Logic', level=1)
        
        self.doc.add_heading('Original Logic:', level=2)
        self.doc.add_paragraph(stage.original_logic or 'N/A')
        
        if stage.edited_logic:
            self.doc.add_heading('Edited Logic:', level=2)
            self.doc.add_paragraph(stage.edited_logic)
        
        # Footer
        self.doc.add_paragraph()
        footer = self.doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save document
        output_dir = Path(settings.BASE_DIR) / "data" / "reports" / "versions"
        output_dir.mkdir(parents=True, exist_ok=True)
        
        filename = f"version_history_{stage.stage_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = output_dir / filename
        
        self.doc.save(str(output_path))
        
        return str(output_path)
    
    def generate_single_version_report(self, stage, version, code, project_name: str) -> str:
        """
        Generate report for a specific version
        
        Args:
            stage: Stage object
            version: VersionHistory object
            code: GeneratedCode object
            project_name: Name of the project
            
        Returns:
            Path to generated DOCX file
        """
        self.doc = Document()
        
        # Title
        title = self.doc.add_heading(f'Version {version.version_number} Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Version Info
        self.doc.add_heading('Version Information', level=1)
        info_table = self.doc.add_table(rows=6, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        metadata = version.version_metadata or {}
        
        info_data = [
            ('Project', project_name),
            ('Stage', f"{stage.stage_number} - {stage.stage_name}"),
            ('Version', version.version_number),
            ('Action', metadata.get('action', 'Unknown')),
            ('Timestamp', version.timestamp.strftime('%Y-%m-%d %H:%M:%S') if version.timestamp else 'N/A'),
            ('Description', metadata.get('description', 'N/A'))
        ]
        
        for i, (key, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = key
            info_table.rows[i].cells[1].text = str(value)
        
        self.doc.add_paragraph()
        
        # Stage Logic
        self.doc.add_heading('Stage Logic', level=1)
        logic = stage.edited_logic if stage.edited_logic else stage.original_logic
        self.doc.add_paragraph(logic or 'N/A')
        
        # Generated Code (if available)
        if code:
            self.doc.add_heading('Generated Code', level=1)
            
            # Global Labels
            self.doc.add_heading('Global Labels:', level=2)
            if code.global_labels:
                labels_table = self.doc.add_table(rows=1, cols=4)
                labels_table.style = 'Light Grid Accent 1'
                
                headers = labels_table.rows[0].cells
                headers[0].text = 'Name'
                headers[1].text = 'Data Type'
                headers[2].text = 'Class'
                headers[3].text = 'Device'
                
                for label in code.global_labels:
                    row = labels_table.add_row().cells
                    row[0].text = label.get('name', '')
                    row[1].text = label.get('data_type', '')
                    row[2].text = label.get('class', '')
                    row[3].text = label.get('device', '')
            else:
                self.doc.add_paragraph('No global labels')
            
            self.doc.add_paragraph()
            
            # Local Labels
            self.doc.add_heading('Local Labels:', level=2)
            if code.local_labels:
                labels_table = self.doc.add_table(rows=1, cols=3)
                labels_table.style = 'Light Grid Accent 1'
                
                headers = labels_table.rows[0].cells
                headers[0].text = 'Name'
                headers[1].text = 'Data Type'
                headers[2].text = 'Class'
                
                for label in code.local_labels:
                    row = labels_table.add_row().cells
                    row[0].text = label.get('name', '')
                    row[1].text = label.get('data_type', '')
                    row[2].text = label.get('class', '')
            else:
                self.doc.add_paragraph('No local labels')
            
            self.doc.add_paragraph()
            
            # Program Body
            self.doc.add_heading('Program Body:', level=2)
            code_para = self.doc.add_paragraph(code.program_body or 'N/A')
            code_para.style = 'Intense Quote'
        
        # Footer
        self.doc.add_paragraph()
        footer = self.doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save document
        output_dir = Path(settings.BASE_DIR) / "data" / "reports" / "versions"
        output_dir.mkdir(parents=True, exist_ok=True)
        
        filename = f"{stage.stage_name}_v{version.version_number}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = output_dir / filename
        
        self.doc.save(str(output_path))
        
        return str(output_path)


# Singleton instance
version_report_generator = VersionReportGenerator()
