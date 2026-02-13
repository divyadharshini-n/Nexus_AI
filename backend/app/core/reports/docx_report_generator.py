from pathlib import Path
from typing import Dict, List
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


class DOCXReportGenerator:
    """Generate professional Word document reports"""
    
    def __init__(self):
        self.doc = None
    
    def generate_project_report(
        self,
        project: Dict,
        stages: List[Dict],
        codes: List[Dict],
        admin_name: str = None,
        validations: Dict = None,
        safety_assessments: Dict = None
    ) -> str:
        """
        Generate complete project report in DOCX format
        
        Returns:
            Path to generated DOCX file
        """
        # Create document
        self.doc = Document()
        
        # Set up styles
        self._setup_styles()
        
        # Add content with all requested sections
        self._add_cover_page(project, admin_name)
        self._add_index_page(project, stages, codes, admin_name)
        self._add_project_overview(project, admin_name, stages)
        self._add_stages_section(stages, codes)
        self._add_generated_code_section(codes)
        self._add_labels_section(codes)
        
        if validations:
            self._add_validation_section(validations)
        
        if safety_assessments:
            self._add_safety_section(safety_assessments)
        
        # Save document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{project['name'].replace(' ', '_')}_{timestamp}.docx"
        output_path = Path("data/reports") / filename
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        self.doc.save(str(output_path))
        
        return str(output_path)
    
    def _setup_styles(self):
        """Set up document styles"""
        styles = self.doc.styles
        
        # Title style
        if 'Report Title' not in [s.name for s in styles]:
            title_style = styles.add_style('Report Title', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Calibri'
            title_style.font.size = Pt(28)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(0, 51, 102)
    
    def _add_cover_page(self, project: Dict, admin_name: str = None):
        """Add professional cover page"""
        # Title
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("AI PLC Platform\n")
        title_run.font.size = Pt(36)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Subtitle
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run("Control Logic Report")
        subtitle_run.font.size = Pt(24)
        subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
        
        # Project name
        self.doc.add_paragraph("\n" * 4)
        project_name = self.doc.add_paragraph()
        project_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = project_name.add_run(project['name'])
        name_run.font.size = Pt(22)
        name_run.font.bold = True
        
        # Admin info
        if admin_name:
            self.doc.add_paragraph("\n" * 2)
            admin_para = self.doc.add_paragraph()
            admin_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            admin_run = admin_para.add_run(f"Created by: {admin_name}")
            admin_run.font.size = Pt(14)
            admin_run.font.italic = True
            admin_run.font.color.rgb = RGBColor(80, 80, 80)
        
        # Date
        self.doc.add_paragraph("\n" * 8)
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        date_run.font.size = Pt(12)
        date_run.font.color.rgb = RGBColor(120, 120, 120)
        
        # Page break
        self.doc.add_page_break()
    
    def _add_index_page(self, project: Dict, stages: List[Dict], codes: List[Dict], admin_name: str = None):
        """Add comprehensive index page with project summary"""
        self.doc.add_heading('Index & Project Summary', 1)
        
        # Project Information Table
        self.doc.add_heading('Project Information', 2)
        info_table = self.doc.add_table(rows=5, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        # Set column widths for better alignment
        for row in info_table.rows:
            row.cells[0].width = Inches(2.5)
            row.cells[1].width = Inches(4.5)
        
        info_table.rows[0].cells[0].text = 'Project Name'
        info_table.rows[0].cells[1].text = project['name']
        
        info_table.rows[1].cells[0].text = 'Admin/Creator'
        info_table.rows[1].cells[1].text = admin_name or 'N/A'
        
        info_table.rows[2].cells[0].text = 'Report Generated'
        info_table.rows[2].cells[1].text = datetime.now().strftime('%B %d, %Y at %I:%M %p')
        
        info_table.rows[3].cells[0].text = 'Total Stages'
        info_table.rows[3].cells[1].text = str(len(stages))
        
        info_table.rows[4].cells[0].text = 'Code Modules Generated'
        info_table.rows[4].cells[1].text = str(len(codes))
        
        self.doc.add_paragraph()
        
        # Stages Summary Table with Validation Status
        self.doc.add_heading('Stages Overview', 2)
        stages_table = self.doc.add_table(rows=1, cols=5)
        stages_table.style = 'Medium Shading 1 Accent 1'
        
        # Headers
        header_cells = stages_table.rows[0].cells
        header_cells[0].text = 'Stage #'
        header_cells[1].text = 'Stage Name'
        header_cells[2].text = 'Type'
        header_cells[3].text = 'Validated'
        header_cells[4].text = 'Finalized'
        
        # Bold headers
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Sort stages by stage_number to ensure correct order
        sorted_stages = sorted(stages, key=lambda x: x.get('stage_number', 0))
        
        # Add stage data rows
        for stage in sorted_stages:
            row_cells = stages_table.add_row().cells
            row_cells[0].text = str(stage.get('stage_number', 'N/A'))
            row_cells[1].text = stage.get('stage_name', 'N/A')
            row_cells[2].text = stage.get('stage_type', 'N/A')
            row_cells[3].text = 'Yes' if stage.get('is_validated') else 'No'
            row_cells[4].text = 'Yes' if stage.get('is_finalized') else 'No'
            
            # Color code validation status
            validation_cell = row_cells[3]
            for paragraph in validation_cell.paragraphs:
                for run in paragraph.runs:
                    if stage.get('is_validated'):
                        run.font.color.rgb = RGBColor(0, 128, 0)
                        run.font.bold = True
                    else:
                        run.font.color.rgb = RGBColor(200, 0, 0)
        
        self.doc.add_paragraph()
        
        # Table of Contents
        self.doc.add_heading('Table of Contents', 2)
        toc_para = self.doc.add_paragraph()
        toc_items = [
            "1. Project Overview",
            "2. Stages & Control Logic",
            "3. Generated Code",
            "4. Global & Local Labels",
            "5. Validation Results",
            "6. Safety Assessments"
        ]
        for item in toc_items:
            toc_para.add_run(f"{item}\n")
        
        self.doc.add_page_break()
    
    def _add_project_overview(self, project: Dict, admin_name: str = None, stages: List[Dict] = None):
        """Add comprehensive project overview section"""
        self.doc.add_heading('1. Project Overview', 1)
        
        # Project Details Table
        table = self.doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Add data with professional alignment
        table.rows[0].cells[0].text = 'Project Name'
        table.rows[0].cells[1].text = project['name']
        
        table.rows[1].cells[0].text = 'Description'
        table.rows[1].cells[1].text = project.get('description', 'N/A')
        
        table.rows[2].cells[0].text = 'Created By (Admin)'
        table.rows[2].cells[1].text = admin_name or 'N/A'
        
        table.rows[3].cells[0].text = 'Created Date'
        table.rows[3].cells[1].text = str(project.get('created_at', 'N/A'))
        
        table.rows[4].cells[0].text = 'Report Generated'
        table.rows[4].cells[1].text = datetime.now().strftime('%B %d, %Y at %I:%M %p')
        
        # Calculate statistics
        if stages:
            validated_count = sum(1 for s in stages if s.get('is_validated'))
            table.rows[5].cells[0].text = 'Validation Status'
            table.rows[5].cells[1].text = f"{validated_count} of {len(stages)} stages validated"
        
        # Bold the first column
        for row in table.rows:
            cell = row.cells[0]
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        self.doc.add_paragraph()
        self.doc.add_page_break()
    
    def _add_stages_section(self, stages: List[Dict], codes: List[Dict]):
        """Add detailed stages section with proper ordering"""
        self.doc.add_heading('2. Stages & Control Logic', 1)
        
        # Sort stages by stage_number for correct display order
        sorted_stages = sorted(stages, key=lambda x: x.get('stage_number', 0))
        
        for stage in sorted_stages:
            # Stage header with number
            self.doc.add_heading(f"Stage {stage['stage_number']}: {stage['stage_name']}", 2)
            
            # Stage information table
            info_table = self.doc.add_table(rows=5, cols=2)
            info_table.style = 'Light List Accent 1'
            
            info_table.rows[0].cells[0].text = 'Stage Type'
            info_table.rows[0].cells[1].text = stage.get('stage_type', 'N/A')
            
            info_table.rows[1].cells[0].text = 'Description'
            info_table.rows[1].cells[1].text = stage.get('description', 'N/A')
            
            info_table.rows[2].cells[0].text = 'Validated'
            validated_text = 'Yes ✓' if stage.get('is_validated') else 'No ✗'
            info_table.rows[2].cells[1].text = validated_text
            # Color code
            if stage.get('is_validated'):
                for p in info_table.rows[2].cells[1].paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(0, 128, 0)
                        r.font.bold = True
            else:
                for p in info_table.rows[2].cells[1].paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(200, 0, 0)
            
            info_table.rows[3].cells[0].text = 'Finalized'
            finalized_text = 'Yes ✓' if stage.get('is_finalized') else 'No ✗'
            info_table.rows[3].cells[1].text = finalized_text
            
            info_table.rows[4].cells[0].text = 'Version Number'
            info_table.rows[4].cells[1].text = str(stage.get('version_number', '1'))
            
            # Bold first column
            for row in info_table.rows:
                for p in row.cells[0].paragraphs:
                    for r in p.runs:
                        r.font.bold = True
            
            self.doc.add_paragraph()
            
            # Control Logic
            self.doc.add_heading('Control Logic:', 3)
            logic_text = stage.get('edited_logic') or stage.get('original_logic') or 'N/A'
            logic_para = self.doc.add_paragraph(logic_text)
            logic_para.style = 'Intense Quote'
            
            self.doc.add_paragraph()  # Spacing between stages
        
        self.doc.add_page_break()
    
    def _add_generated_code_section(self, codes: List[Dict]):
        """Add section with last generated code for each stage"""
        self.doc.add_heading('3. Generated Code', 1)
        
        if not codes:
            self.doc.add_paragraph("No code has been generated yet.")
            self.doc.add_page_break()
            return
        
        for idx, code in enumerate(codes, 1):
            # Program header
            program_name = code.get('program_name', f'Program_{idx}')
            self.doc.add_heading(f"Program: {program_name}", 2)
            
            # Program details
            details_table = self.doc.add_table(rows=2, cols=2)
            details_table.style = 'Light List Accent 1'
            
            details_table.rows[0].cells[0].text = 'Execution Type'
            details_table.rows[0].cells[1].text = code.get('execution_type', 'N/A')
            
            details_table.rows[1].cells[0].text = 'Program Name'
            details_table.rows[1].cells[1].text = program_name
            
            # Bold first column
            for row in details_table.rows:
                for p in row.cells[0].paragraphs:
                    for r in p.runs:
                        r.font.bold = True
            
            self.doc.add_paragraph()
            
            # Program Body - The actual generated code
            if code.get('program_body'):
                self.doc.add_heading('Structured Text Code:', 3)
                code_para = self.doc.add_paragraph(code['program_body'])
                code_para.style = 'Intense Quote'
                
                # Format as code
                for run in code_para.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
            
            self.doc.add_paragraph()  # Spacing
        
        self.doc.add_page_break()
    
    def _add_labels_section(self, codes: List[Dict]):
        """Add comprehensive labels section showing all global and local labels"""
        self.doc.add_heading('4. Global & Local Labels', 1)
        
        if not codes:
            self.doc.add_paragraph("No labels available.")
            self.doc.add_page_break()
            return
        
        for idx, code in enumerate(codes, 1):
            program_name = code.get('program_name', f'Program_{idx}')
            self.doc.add_heading(f"{program_name} - Labels", 2)
            
            # Global Labels Section
            self.doc.add_heading('Global Labels:', 3)
            global_labels = code.get('global_labels', [])
            
            if global_labels:
                # Create table with proper columns
                global_table = self.doc.add_table(rows=1, cols=6)
                global_table.style = 'Medium Shading 1 Accent 1'
                
                # Header
                header_cells = global_table.rows[0].cells
                headers = ['Name', 'Data Type', 'Class', 'Device', 'Initial Value', 'Comment']
                for i, header in enumerate(headers):
                    header_cells[i].text = header
                    for p in header_cells[i].paragraphs:
                        for r in p.runs:
                            r.font.bold = True
                            r.font.color.rgb = RGBColor(255, 255, 255)
                
                # Data rows
                for label in global_labels:
                    row_cells = global_table.add_row().cells
                    row_cells[0].text = label.get('name', '')
                    row_cells[1].text = label.get('data_type', '')
                    row_cells[2].text = label.get('class', '')
                    row_cells[3].text = label.get('device', '')
                    row_cells[4].text = str(label.get('initial_value', ''))
                    row_cells[5].text = label.get('comment', '')
            else:
                self.doc.add_paragraph("No global labels defined.")
            
            self.doc.add_paragraph()
            
            # Local Labels Section
            self.doc.add_heading('Local Labels:', 3)
            local_labels = code.get('local_labels', [])
            
            if local_labels:
                # Create table
                local_table = self.doc.add_table(rows=1, cols=5)
                local_table.style = 'Medium Shading 1 Accent 1'
                
                # Header
                header_cells = local_table.rows[0].cells
                headers = ['Name', 'Data Type', 'Class', 'Initial Value', 'Comment']
                for i, header in enumerate(headers):
                    header_cells[i].text = header
                    for p in header_cells[i].paragraphs:
                        for r in p.runs:
                            r.font.bold = True
                            r.font.color.rgb = RGBColor(255, 255, 255)
                
                # Data rows
                for label in local_labels:
                    row_cells = local_table.add_row().cells
                    row_cells[0].text = label.get('name', '')
                    row_cells[1].text = label.get('data_type', '')
                    row_cells[2].text = label.get('class', '')
                    row_cells[3].text = str(label.get('initial_value', ''))
                    row_cells[4].text = label.get('comment', '')
            else:
                self.doc.add_paragraph("No local labels defined.")
            
            self.doc.add_paragraph()  # Spacing between programs
        
        self.doc.add_page_break()
    
    def _add_validation_section(self, validations: Dict):
        """Add comprehensive validation results section"""
        self.doc.add_heading('5. Validation Results', 1)
        
        if not validations:
            self.doc.add_paragraph("No validation results available.")
            self.doc.add_page_break()
            return
        
        for stage_id, validation in validations.items():
            self.doc.add_heading(f"Stage {stage_id} Validation", 2)
            
            # Status with color coding
            status_para = self.doc.add_paragraph()
            status_text = f"Status: {validation.get('status', 'N/A')}"
            status_run = status_para.add_run(status_text)
            status_run.bold = True
            status_run.font.size = Pt(12)
            
            if validation.get('status') == 'PASS':
                status_run.font.color.rgb = RGBColor(0, 128, 0)
            elif validation.get('status') == 'FAIL':
                status_run.font.color.rgb = RGBColor(255, 0, 0)
            else:
                status_run.font.color.rgb = RGBColor(255, 165, 0)
            
            # Analysis sections
            if validation.get('semantic_analysis'):
                self.doc.add_heading('Semantic Analysis:', 3)
                self.doc.add_paragraph(validation['semantic_analysis'])
            
            if validation.get('logical_consistency'):
                self.doc.add_heading('Logical Consistency:', 3)
                self.doc.add_paragraph(validation['logical_consistency'])
            
            if validation.get('safety_compliance'):
                self.doc.add_heading('Safety Compliance:', 3)
                self.doc.add_paragraph(validation['safety_compliance'])
            
            # Issues
            if validation.get('issues'):
                self.doc.add_heading('Issues Found:', 3)
                for issue in validation['issues']:
                    p = self.doc.add_paragraph(f"⚠ {issue}", style='List Bullet')
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(200, 0, 0)
            
            # Recommendations
            if validation.get('recommendations'):
                self.doc.add_heading('Recommendations:', 3)
                for rec in validation['recommendations']:
                    p = self.doc.add_paragraph(f"💡 {rec}", style='List Bullet')
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(0, 100, 200)
            
            self.doc.add_paragraph()  # Spacing
        
        self.doc.add_page_break()
    
    def _add_safety_section(self, safety_assessments: Dict):
        """Add comprehensive safety assessment section"""
        self.doc.add_heading('6. Safety Assessments', 1)
        
        if not safety_assessments:
            self.doc.add_paragraph("No safety assessments available.")
            return
        
        for stage_id, assessment in safety_assessments.items():
            self.doc.add_heading(f"Stage {stage_id} Safety Assessment", 2)
            
            # Status
            status_para = self.doc.add_paragraph()
            status_run = status_para.add_run(f"Safety Status: {assessment.get('status', 'N/A')}")
            status_run.bold = True
            status_run.font.size = Pt(12)
            
            # Hazards
            if assessment.get('hazards'):
                self.doc.add_heading('Potential Hazards:', 3)
                for hazard in assessment['hazards']:
                    p = self.doc.add_paragraph(f"⚠️ {hazard}", style='List Bullet')
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(255, 140, 0)
            
            # Violations
            if assessment.get('violations'):
                self.doc.add_heading('Safety Violations:', 3)
                for violation in assessment['violations']:
                    p = self.doc.add_paragraph(f"🚨 {violation}", style='List Bullet')
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(200, 0, 0)
            
            # Mitigation measures
            if assessment.get('mitigations'):
                self.doc.add_heading('Mitigation Measures:', 3)
                for mitigation in assessment['mitigations']:
                    p = self.doc.add_paragraph(f"✓ {mitigation}", style='List Bullet')
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(0, 128, 0)
            
            self.doc.add_paragraph()  # Spacing


# Global instance
docx_report_generator = DOCXReportGenerator()